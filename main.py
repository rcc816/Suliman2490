from flet import *
import sqlite3
import os
from datetime import datetime
# from fpdf import FPDF
from collections import Counter
import csv
from lang import tr, lang
import lang as lang_module
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
# import arabic_reshaper
# from bidi.algorithm import get_display

def reshape_ar(text):
    return get_display(arabic_reshaper.reshape(text))

lang_module.lang = "ar"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
db_path = os.path.join(BASE_DIR, "db", "diseases.db")
logo_path = os.path.join(os.path.dirname(__file__), "logo.png")

def msg(page, text, lang="ar"):
    dlgmsg = AlertDialog(
        title=Text(tr("notice", lang)),
        content=Text(text),
        alignment=alignment.center,
        on_dismiss=lambda e: None,
        title_padding=padding.all(25),
    )
    page.open(dlgmsg)
    page.update()

def about(page):
    about_text_lines =[
            "مشروع تخرج الطالبة / هاجر بنت سليمان الرشيد",
            "كلية العلوم الطبية التطبيقية ","تخصص معلوماتية صحية",
            "العام الدراسي 1447 هـ 2025-2026 م"
    ]
    dlgmsg = AlertDialog(
        title=Text("About حول"),
        content=Column(
            [Text(line) for line in about_text_lines],
            horizontal_alignment=CrossAxisAlignment.CENTER
            ),
        actions=[
            TextButton("موافق", on_click=lambda _: page.close(dlgmsg)),
        ],
        icon=Icon(Icons.INFO_ROUNDED, size=40, color=Colors.GREEN_ACCENT_700),
        alignment=alignment.center,
        on_dismiss=lambda e: None,
        title_padding=padding.all(25),
    )
    page.open(dlgmsg)
    page.update()    

def error(page, text, lang="ar"):
    dlgerror = AlertDialog(
        title=Text(tr("error", lang), color="red"),
        content=Text(text),
        alignment=alignment.center,
        on_dismiss=lambda e: None,
        title_padding=padding.all(25)
    )
    page.open(dlgerror)
    page.update()

def confirm(page, text, yes_action, lang="ar"):
    def close_dialog(e):
        dlgconfirm.open = False
        page.update()

    def yes(e):
        close_dialog(e)
        yes_action()

    dlgconfirm = AlertDialog(
        modal=True,
        title=Text(tr("confirm", lang)),
        content=Text(text),
        actions=[
            TextButton(tr("yes", lang), on_click=yes),
            TextButton(tr("cancel", lang), on_click=close_dialog),
        ],
    )
    page.open(dlgconfirm)
    page.update()


# ========== قاعدة البيانات ==========
def init_db():
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS diseases (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT,
                    description TEXT
                )""")
    c.execute("""CREATE TABLE IF NOT EXISTS symptoms (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT
                )""")
    c.execute("""CREATE TABLE IF NOT EXISTS disease_symptoms (
                    disease_id INTEGER,
                    symptom_id INTEGER
                )""")
    c.execute("""CREATE TABLE IF NOT EXISTS diagnosis_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    symptom TEXT,
                    result TEXT,
                    date TEXT
                )""")

    conn.commit()
    conn.close()

# ========== الشاشة الرئيسية ==========

def SplashScreen(page, set_lang, go_to_main_menu):
    return Container(
        expand=True,
        alignment=alignment.center,
        content=Column(
            [
                Image(
                    src="logo.png",
                    width=150,
                    height=150,
                ),
                Text("المساعد في تشخيص الأمراض", size=24, weight="bold", text_align="center"),
                Text("Diagnosis Assistant", size=18, italic=True, text_align="center"),
                
                # فراغات تكتب فيها معلومات لاحقًا
                Text("", height=20),
                Text("", height=20),
                Text("", height=20),

                Row(
                    alignment="center",
                    controls=[
                        ElevatedButton("العربية", on_click=lambda e: (set_lang("ar"), go_to_main_menu())),
                        ElevatedButton("English", on_click=lambda e: (set_lang("en"), go_to_main_menu())),
                    ],
                ),
            ],
            horizontal_alignment="center",
            spacing=15,
            alignment="center"
        ),
    )

def MainMenuScreen(page, lang, go_to_screen):
    def build_button(icon, label, action):
        return ElevatedButton(
            content=Column([
                Icon(icon, size=36),
                Text(label, size=14, text_align="center")
            ], alignment="center"),
            on_click=action,
            style=ButtonStyle(padding=padding.all(15))
        )

    unified_width = 600

    return Column(
        expand=True,
        alignment="center",
        horizontal_alignment="center",
        controls=[
            Container(
                width=unified_width,
                alignment=alignment.center,
                content=Text(tr("main_title", lang), size=26, weight="bold", text_align="center"),
                border=border.all(2, Colors.GREY),
                padding=padding.symmetric(vertical=12, horizontal=20),
                border_radius=12,
                margin=margin.only(bottom=20),
                bgcolor=Colors.BLUE_100
            ),
            Container(
                width=unified_width,
                alignment=alignment.center,
                border=border.all(2, Colors.GREY),
                border_radius=12,
                padding=padding.all(20),
                content=Column([
                    Row(
                        alignment="center",
                        spacing=20,
                        controls=[
                            build_button(Icons.HEALTH_AND_SAFETY, tr("start_diagnosis", lang), lambda e: go_to_screen("diagnosis")),
                            build_button(Icons.ADD_BOX, tr("add_disease", lang), lambda e: go_to_screen("add_disease")),
                            build_button(Icons.LIST_ALT, tr("add_symptom", lang), lambda e: go_to_screen("manage_symptoms"))
                        ]
                    ),
                    Row(
                        alignment="center",
                        spacing=20,
                        controls=[
                            build_button(Icons.HISTORY, tr("diagnosis_history", lang), lambda e: go_to_screen("history")),
                            build_button(Icons.INSIGHTS, tr("data_analysis", lang), lambda e: go_to_screen("data_analysis")),
                            build_button(Icons.UPLOAD, tr("export_data", lang), lambda e: go_to_screen("data_export_import"))
                        ]
                    )
                ]),
                bgcolor=Colors.GREY_100
            ),
            Container(
                width=unified_width,
                alignment=alignment.center,
                content=ElevatedButton(
                    content=Column([
                        Icon(Icons.EXIT_TO_APP, size=36),
                        Text(tr("exit", lang), size=14, text_align="center")
                    ], alignment="center"),
                    on_click=lambda e: page.window.close(),
                    style=ButtonStyle(padding=padding.all(15))
                ),
                border=border.all(2, Colors.GREY),
                border_radius=12,
                padding=padding.symmetric(vertical=10, horizontal=30),
                margin=margin.only(top=20),
                bgcolor=Colors.BLUE_100
            )
        ]
    )

# ========== شاشة التشخيص ==========
def DiagnosisScreen(page, go_back):
    symptoms_selected = []
    rtl = lang_module.lang == "ar"

    symptom_input = TextField(label=tr("select_symptoms"), expand=True, text_align="right" if rtl else "left")
    symptom_dropdown = Dropdown(label=tr("select_symptoms"), options=[], width=300)
    result_text = Text(tr("results") + ":", size=16, weight="bold", text_align="right" if rtl else "left")
    result_display = Text("", selectable=True, text_align="right" if rtl else "left")
    selected_symptoms = Row(wrap=True, spacing=5, alignment="end" if rtl else "start")

    def load_symptoms():
        symptom_dropdown.options.clear()
        conn = sqlite3.connect("db/diseases.db")
        c = conn.cursor()
        column = "name_ar" if lang_module.lang == "ar" else "name_en"
        c.execute(f"SELECT {column} FROM symptoms")
        for row in c.fetchall():
            symptom_dropdown.options.append(dropdown.Option(row[0]))
        conn.close()
        page.update()

    def add_symptom(e):
        selected = symptom_dropdown.value or symptom_input.value.strip()
        if selected and selected not in symptoms_selected:
            symptoms_selected.append(selected)
            selected_symptoms.controls.append(
                Chip(label=Text(selected), on_delete=lambda _: remove_symptom(selected))
            )
            page.update()

    def remove_symptom(symptom):
        symptoms_selected.remove(symptom)
        selected_symptoms.controls.clear()
        for s in symptoms_selected:
            selected_symptoms.controls.append(
                Chip(label=Text(s), on_delete=lambda _: remove_symptom(s))
            )
        page.update()

    def diagnose(e):
        if not symptoms_selected:
            error(page, tr("no_symptoms_selected"), lang_module.lang)
            return

        conn = sqlite3.connect("db/diseases.db")
        c = conn.cursor()

        placeholders = ",".join("?" for _ in symptoms_selected)
        lang_col = "name_ar" if lang_module.lang == "ar" else "name_en"
        c.execute(f"SELECT id FROM symptoms WHERE {lang_col} IN ({placeholders})", symptoms_selected)
        symptom_ids = [row[0] for row in c.fetchall()]

        if not symptom_ids:
            result_display.value = tr("no_diseases_found", lang_module.lang)
            page.update()
            return

        placeholders = ",".join("?" for _ in symptom_ids)
        c.execute(f"""
            SELECT disease_id FROM disease_symptoms
            WHERE symptom_id IN ({placeholders})
            GROUP BY disease_id
            HAVING COUNT(*) >= 1
        """, symptom_ids)

        disease_ids = [row[0] for row in c.fetchall()]

        if not disease_ids:
            result_display.value = tr("no_diseases_found", lang_module.lang)
            page.update()
            return

        placeholders = ",".join("?" for _ in disease_ids)
        name_column = "name" if lang_module.lang == "ar" else "name_en"
        c.execute(f"SELECT {name_column} FROM diseases WHERE id IN ({placeholders})", disease_ids)
        diseases = [row[0] for row in c.fetchall()]

        if diseases:
            result_display.value = "\n".join(diseases)
        else:
            result_display.value = tr("no_diseases_found", lang_module.lang)

        c.execute(
         "INSERT INTO diagnosis_history (symptoms, result, diagnosis_date) VALUES (?, ?, ?)",
        (", ".join(symptoms_selected), result_display.value, datetime.now().strftime("%Y-%m-%d %H:%M"))
        )

        conn.commit()
        conn.close()
        page.update()

    def set_paragraph_rtl(paragraph):
        paragraph.paragraph_format.alignment = 2  # Align Right
        paragraph.paragraph_format.right_indent = Cm(1)
        paragraph.paragraph_format.first_line_indent = Cm(0.5)
        paragraph.paragraph_format.space_after = Pt(8)
        
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def export_to_word(symptoms_selected, result_text, page):
        document = Document()

        # إعداد الخط
        style = document.styles['Normal']
        font = style.font
        font.name = 'Traditional Arabic' if lang_module.lang == "ar" else 'Calibri'
        font.size = Pt(14)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font.name)

        # إدراج اللوقو
        logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
        try:
            if os.path.exists(logo_path):
                document.add_picture(logo_path, width=Inches(2))
        except Exception as e:
            print("لم يتم إدراج اللوقو:", e)

        # العنوان
        heading = document.add_heading(tr("report_title", lang_module.lang), level=1)
        if lang_module.lang == "ar":
            set_paragraph_rtl(heading)

        # التاريخ
        p = document.add_paragraph(f"{tr('date', lang_module.lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if lang_module.lang == "ar":
            set_paragraph_rtl(p)

        # الأعراض
        p = document.add_paragraph(f"{tr('selected_symptoms', lang_module.lang)}: {', '.join(symptoms_selected)}")
        if lang_module.lang == "ar":
            set_paragraph_rtl(p)

        # النتائج
        p = document.add_paragraph(tr("results", lang_module.lang) + ":")
        if lang_module.lang == "ar":
            set_paragraph_rtl(p)

        for r in result_text.splitlines():
            if r.strip():
                bullet = document.add_paragraph(f"• {r.strip()}")
                if lang_module.lang == "ar":
                    set_paragraph_rtl(bullet)

        filename = f"diagnosis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(filename)

        msg(page, f"{tr('report_saved_as', lang_module.lang)}\n{filename}", lang_module.lang)
    load_symptoms()

    return Column(
        scroll=ScrollMode.AUTO,
        horizontal_alignment="center",
        controls=[
            Container(
                width=600,
                alignment=alignment.center,
                content=Row([
                    IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back()),
                    Text(tr("select_symptoms", lang_module.lang), size=20, weight="bold", text_align="right" if rtl else "left")
                ], alignment="end" if rtl else "start"),
                border=border.all(2, Colors.GREY),
                padding=padding.all(12),
                border_radius=12,
                margin=margin.only(bottom=20)
            ),
            Container(
                width=600,
                border=border.all(2, Colors.GREY),
                border_radius=12,
                padding=padding.all(15),
                content=Column([
                    Row([symptom_input, symptom_dropdown], alignment="end" if rtl else "start", spacing=10),
                    ElevatedButton(tr("add_symptom", lang_module.lang), on_click=add_symptom)
                ])
            ),
            Container(
                width=600,
                border=border.all(2, Colors.GREY),
                border_radius=12,
                padding=padding.all(15),
                margin=margin.only(top=20),
                content=Column([
                    selected_symptoms,
                    ElevatedButton(tr("diagnose", lang_module.lang), on_click=diagnose, icon=Icons.SEARCH)
                ])
            ),
            Container(
                width=600,
                border=border.all(2, Colors.GREY),
                border_radius=12,
                padding=padding.all(15),
                margin=margin.only(top=20),
                content=Column([
                    result_text,
                    result_display,
                    ElevatedButton(tr("export_word", lang_module.lang), icon=Icons.DESCRIPTION, on_click=lambda e: export_to_word(symptoms_selected, result_display.value, page))
                ])
            )
        ]
    )




def AddDiseaseScreen(page, go_back):
    name_ar_field = TextField(label=tr("اسم المرض بالعربية"), expand=True)
    name_en_field = TextField(label=tr("Disease name in English"), expand=True)
    symptoms_options = Column(spacing=5, scroll=ScrollMode.AUTO)
    selected_symptoms = []

    def load_symptoms():
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT name_ar FROM symptoms")
        for row in c.fetchall():
            symptom = row[0]
            chk = Checkbox(label=symptom, on_change=lambda e, s=symptom: toggle_symptom(s, e.control.value))
            symptoms_options.controls.append(chk)
        conn.close()
        page.update()

    def toggle_symptom(symptom, selected):
        if selected:
            if symptom not in selected_symptoms:
                selected_symptoms.append(symptom)
        else:
            if symptom in selected_symptoms:
                selected_symptoms.remove(symptom)

    def save_disease(e):
        name_ar = name_ar_field.value.strip()
        name_en = name_en_field.value.strip()

        if not name_ar or not name_en:
            page.dialog = AlertDialog(title=Text(tr("خطأ")), content=Text(tr("يرجى تعبئة كلا الحقلين")))
            page.dialog.open = True
            page.update()
            return

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("INSERT INTO diseases (name_ar, name_en) VALUES (?, ?)", (name_ar, name_en))
        disease_id = c.lastrowid

        for symptom in selected_symptoms:
            c.execute("INSERT INTO disease_symptoms (disease_id, symptom_name) VALUES (?, ?)", (disease_id, symptom))

        conn.commit()
        conn.close()

        name_ar_field.value = ""
        name_en_field.value = ""
        selected_symptoms.clear()
        for chk in symptoms_options.controls:
            chk.value = False

        page.snack_bar = SnackBar(Text(tr("تمت إضافة المرض بنجاح")), bgcolor=Colors.GREEN)
        page.snack_bar.open = True
        page.update()

    # تحميل الأعراض عند عرض الشاشة
    load_symptoms()

    return Column(
        scroll=ScrollMode.AUTO,
        controls=[
            Row([IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back())]),
            Text(tr("add_disease"), size=24, weight="bold"),
            name_ar_field,
            name_en_field,
            Divider(),
            Text(tr("select_symptoms"), size=18),
            symptoms_options,
            ElevatedButton(tr("Save"), icon=Icons.SAVE, on_click=save_disease)
        ]
    )

def ManageSymptomsScreen(page, go_back):
    name_ar_field = TextField(label=tr("اسم العرض بالعربية"), expand=True)
    name_en_field = TextField(label=tr("Symptom name in English"), expand=True)
    symptoms_list = Column(scroll=ScrollMode.AUTO)

    def load_symptoms():
        symptoms_list.controls.clear()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT id, name_ar, name_en FROM symptoms ORDER BY id DESC")
        for row in c.fetchall():
            s_id, name_ar, name_en = row
            symptoms_list.controls.append(
                Card(
                    content=ListTile(
                        title=Text(f"{name_ar} / {name_en}"),
                        trailing=IconButton(Icons.DELETE, icon_color=Colors.RED, tooltip=tr("delete"), on_click=lambda e, s_id=s_id: delete_symptom(s_id))
                    )
                )
            )
        conn.close()
        page.update()

    def delete_symptom(s_id):
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("DELETE FROM symptoms WHERE id=?", (s_id,))
        c.execute("DELETE FROM disease_symptoms WHERE symptom_name IN (SELECT name_ar FROM symptoms WHERE id=?)", (s_id,))
        conn.commit()
        conn.close()
        load_symptoms()

    def add_symptom(e):
        name_ar = name_ar_field.value.strip()
        name_en = name_en_field.value.strip()

        if not name_ar or not name_en:
            page.dialog = AlertDialog(title=Text(tr("خطأ")), content=Text(tr("يرجى تعبئة كلا الحقلين")))
            page.dialog.open = True
            page.update()
            return

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("INSERT INTO symptoms (name_ar, name_en) VALUES (?, ?)", (name_ar, name_en))
        conn.commit()
        conn.close()

        name_ar_field.value = ""
        name_en_field.value = ""

        page.snack_bar = SnackBar(Text(tr("تمت إضافة العرض بنجاح")), bgcolor=Colors.GREEN)
        page.snack_bar.open = True
        load_symptoms()

    # تحميل الأعراض عند بداية العرض
    load_symptoms()

    return Column(
        scroll=ScrollMode.AUTO,
        controls=[
            Row([IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back())]),
            Text(tr("add_symptom"), size=24, weight="bold"),
            name_ar_field,
            name_en_field,
            ElevatedButton(tr("Save"), icon=Icons.SAVE, on_click=add_symptom),
            Divider(),
            Text(tr("add_symptom") + ":", size=18),
            symptoms_list
        ]
    )


def link_symptoms_screen(page, lang):
    disease_dropdown = Dropdown(label=tr("اختر مرضًا", "Select Disease", lang), width=300)
    symptom_checkboxes = Column(scroll=ScrollMode.AUTO)
    linked_symptoms_list = Column(scroll=ScrollMode.AUTO)

    def load_diseases():
        disease_dropdown.options.clear()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT id, name FROM diseases")
        for row in c.fetchall():
            disease_dropdown.options.append(dropdowns.Option(key=row[0], text=row[1]))
        conn.close()
        page.update()

    def load_symptoms():
        symptom_checkboxes.controls.clear()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT id, name FROM symptoms")
        for row in c.fetchall():
            checkbox = Checkbox(label=row[1], data=row[0])
            symptom_checkboxes.controls.append(checkbox)
        conn.close()
        page.update()

    def load_linked_symptoms(disease_id):
        linked_symptoms_list.controls.clear()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("""SELECT s.id, s.name FROM symptoms s
                     JOIN disease_symptoms ds ON s.id = ds.symptom_id
                     WHERE ds.disease_id = ?""", (disease_id,))
        for sid, name in c.fetchall():
            def delete_handler(e, s_id=sid):
                confirm(page, tr("هل تريد حذف الارتباط؟", "Delete this link?", lang), lambda: unlink_symptom(disease_id, s_id))

            linked_symptoms_list.controls.append(
                Container(
                    content=Row([
                        Text(name, expand=True),
                        IconButton(Icons.DELETE, icon_color="red", on_click=delete_handler)
                    ]),
                    border=border.all(1, Colors.GREY),
                    padding=8,
                    margin=5
                )
            )
        conn.close()
        page.update()

    def link_symptoms(e):
        disease_id = disease_dropdown.value
        if not disease_id:
            error(page, tr("اختر مرضًا أولاً", "Please select a disease", lang), lang)
            return

        selected_ids = [cb.data for cb in symptom_checkboxes.controls if cb.value]
        if not selected_ids:
            error(page, tr("حدد عرضًا واحدًا على الأقل", "Select at least one symptom", lang), lang)
            return

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        for sid in selected_ids:
            c.execute("INSERT OR IGNORE INTO disease_symptoms (disease_id, symptom_id) VALUES (?, ?)", (disease_id, sid))
        conn.commit()
        conn.close()
        load_linked_symptoms(disease_id)
        page.snack_bar = SnackBar(Text(tr("تم الربط بنجاح", "Linked successfully", lang)))
        page.snack_bar.open = True
        page.update()

    def on_disease_change(e):
        if disease_dropdown.value:
            load_linked_symptoms(disease_dropdown.value)

    def unlink_symptom(did, sid):
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("DELETE FROM disease_symptoms WHERE disease_id = ? AND symptom_id = ?", (did, sid))
        conn.commit()
        conn.close()
        load_linked_symptoms(did)

    load_diseases()
    load_symptoms()

    return View(
        "/linking",
        controls=[
            AppBar(title=Text(tr("ربط الأعراض بالأمراض", "Link Symptoms to Diseases", lang)), leading=IconButton(Icons.ARROW_BACK, on_click=lambda e: page.go("/"))),
            Column([
                disease_dropdown,
                ElevatedButton(tr("تحميل الأعراض المرتبطة", "Load Linked Symptoms", lang), on_click=on_disease_change),
                Divider(),
                Row([
                    Column([
                        Text(tr("جميع الأعراض", "All Symptoms", lang), weight="bold"),
                        symptom_checkboxes,
                        ElevatedButton(tr("ربط الأعراض", "Link Symptoms", lang), on_click=link_symptoms)
                    ], expand=True),
                    Column([
                        Text(tr("الأعراض المرتبطة", "Linked Symptoms", lang), weight="bold"),
                        linked_symptoms_list
                    ], expand=True),
                ], spacing=20)
            ])
        ]
    )

def DiagnosisHistoryScreen(page, go_back):
    history_list = Column(scroll=ScrollMode.AUTO)

    def set_rtl(paragraph):
        paragraph.alignment = 2
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def export_all_records(page, lang="ar"):
        try:
            conn = sqlite3.connect("db/diseases.db")
            c = conn.cursor()
            c.execute("SELECT symptoms, result, diagnosis_date FROM diagnosis_history ORDER BY id DESC")
            records = c.fetchall()
            conn.close()

            if not records:
                error(page, "لا توجد بيانات لتصديرها" if lang == "ar" else "No records to export")
                return

            doc = Document()

            # إدراج اللوقو في الأعلى
            logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
            if os.path.exists(logo_path):
                doc.add_picture(logo_path, width=Inches(1.5))
                doc.paragraphs[-1].alignment = 1

            # العنوان
            if lang == "ar":
                heading = doc.add_heading("سجل التشخيص الكامل", level=1)
            else:
                heading = doc.add_heading("Full Diagnosis Report", level=1)
            heading.alignment = 1

            for idx, (symptoms, result, date) in enumerate(records, start=1):
                if lang == "ar":
                    p1 = doc.add_paragraph(f"\nتشخيص رقم {idx}")
                    set_rtl(p1)
                    p2 = doc.add_paragraph(f"التاريخ: {date}")
                    set_rtl(p2)
                    p3 = doc.add_paragraph(f"الأعراض: {symptoms}")
                    set_rtl(p3)
                    p4 = doc.add_paragraph(f"النتيجة: {result}")
                    set_rtl(p4)
                else:
                    doc.add_paragraph(f"\nDiagnosis #{idx}", style="List Number")
                    doc.add_paragraph(f"Date: {date}")
                    doc.add_paragraph(f"Symptoms: {symptoms}")
                    doc.add_paragraph(f"Result: {result}")

            os.makedirs("exports", exist_ok=True)
            filename = f"diagnosis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("exports", filename)
            doc.save(filepath)

            msg(page, "تم تصدير السجل الكامل بنجاح إلى ملف Word" if lang == "ar" else "All records exported to Word successfully")

        except Exception as e:
            error(page, f"فشل التصدير: {str(e)}" if lang == "ar" else f"Export failed: {str(e)}")

    def delete_record(h_id):
        confirm(
            page,
            "هل أنت متأكد من حذف هذا التشخيص؟",
            yes_action=lambda: confirm_and_delete(h_id)
        )

    def confirm_and_delete(h_id):
        conn = sqlite3.connect("db/diseases.db")
        c = conn.cursor()
        c.execute("DELETE FROM diagnosis_history WHERE id=?", (h_id,))
        conn.commit()
        conn.close()
        msg(page, "تم حذف التشخيص بنجاح")
        load_history()

    def delete_all_records(e):
        confirm(
            page,
            "هل تريد حذف كل سجل التشخيص؟",
            yes_action=confirm_and_delete_all
        )

    def confirm_and_delete_all():
        conn = sqlite3.connect("db/diseases.db")
        c = conn.cursor()
        c.execute("DELETE FROM diagnosis_history")
        conn.commit()
        conn.close()
        msg(page, "تم حذف كل سجل التشخيص")
        load_history()

    def load_history():
        history_list.controls.clear()
        conn = sqlite3.connect("db/diseases.db")
        c = conn.cursor()
        c.execute("SELECT id, symptoms, result, diagnosis_date FROM diagnosis_history ORDER BY id DESC")
        for row in c.fetchall():
            h_id, symptoms, result, date = row
            history_list.controls.append(
                Card(
                    content=Container(
                        padding=10,
                        border=border.all(1, Colors.GREY),
                        content=Column([
                            Row([
                                Text(f"{tr('selected_symptoms')}: ", weight="bold"),
                                Text(symptoms, expand=True),
                                Text(date, size=12)
                            ]),
                            Text(f"{tr('results')}: {result}", selectable=True),
                            Row([
                                ElevatedButton(
                                    tr("delete"),
                                    icon=Icons.DELETE,
                                    on_click=lambda e, i=h_id: delete_record(i),
                                    bgcolor=Colors.RED
                                )
                            ])
                        ])
                    )
                )
            )
        conn.close()
        page.update()

    load_history()

    return Column(
        scroll=ScrollMode.AUTO,
        controls=[
            Row([
                IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back()),
                Text(tr("diagnosis_history"), size=24, weight="bold"),
                ElevatedButton(tr("delete_all"), icon=Icons.DELETE_FOREVER, bgcolor=Colors.RED_400, on_click=delete_all_records),
            ElevatedButton(tr("export_all"), icon=Icons.DESCRIPTION, on_click=lambda e: export_all_records(page, lang_module.lang)),
            ]),
            history_list
        ]
    )

def DataAnalysisScreen(page, go_back):
    chart = Column()
    message = Text("")

    def load_chart():
        chart.controls.clear()
        message.value = ""

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("SELECT result FROM diagnosis_history")
        results = c.fetchall()
        conn.close()

        # استخراج الأمراض من النصوص (ممكن تكون أكثر من مرض مفصول بسطر)
        all_diseases = []
        for row in results:
            for d in row[0].split("\n"):
                name = d.strip()
                if name:
                    all_diseases.append(name)

        if not all_diseases:
            message.value = tr("لا توجد بيانات")
            page.update()
            return

        disease_counts = Counter(all_diseases).most_common(10)  # أكثر 10 أمراض فقط

        max_count = max([count for _, count in disease_counts])
        for disease, count in disease_counts:
            chart.controls.append(
                Row([
                    Container(width=(count / max_count) * 250, height=24, bgcolor=Colors.BLUE, border_radius=6),
                    Text(f"{count}", size=12, weight="bold"),
                    Text(disease, expand=True)
                ], spacing=10, vertical_alignment="center")
            )

        page.update()

    load_chart()

    return Column(
        scroll=ScrollMode.AUTO,
        controls=[
            Row([IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back())]),
            Text(tr("📊 تقرير التحليل البياني"), size=24, weight="bold"),
            message,
            chart
        ]
    )

from openpyxl import Workbook, load_workbook

def ExportImportScreen(page, go_back):
    def export_data(e):
        try:
            conn = sqlite3.connect(db_path)
            c = conn.cursor()

            filename = f"medical_data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            wb = Workbook()

            tables = [
                ("diseases", "Diseases"),
                ("symptoms", "Symptoms"),
                ("disease_symptoms", "Links")
            ]

            for index, (table, sheet_name) in enumerate(tables):
                if index == 0:
                    ws = wb.active
                    ws.title = sheet_name
                else:
                    ws = wb.create_sheet(title=sheet_name)

                c.execute(f"SELECT * FROM {table}")
                rows = c.fetchall()
                headers = [desc[0] for desc in c.description]

                ws.append(headers)
                for row in rows:
                    ws.append(row)

            conn.close()
            wb.save(filename)

            msg(page, tr("report_saved", lang_module.lang))

        except Exception as ex:
            print("Export Error:", ex)
            error(page, tr("فشل التصدير", lang_module.lang))

    def import_data(e):
        try:
            file_picker = FilePicker()
            page.overlay.append(file_picker)
            page.update()

            def on_file_selected(ev: FilePickerResultEvent):
                if ev.files:
                    filepath = ev.files[0].path
                    wb = load_workbook(filepath)

                    conn = sqlite3.connect(db_path)
                    c = conn.cursor()

                    # قائمة الجداول المتوقع وجودها في الملف
                    sheets = {
                        "Diseases": ("diseases", ["id", "name", "name_en"]),
                        "Symptoms": ("symptoms", ["id", "name_ar", "name_en"]),
                        "Links": ("disease_symptoms", ["disease_id", "symptom_id"])
                    }

                    for sheet_name, (table_name, columns) in sheets.items():
                        if sheet_name in wb.sheetnames:
                            ws = wb[sheet_name]
                            rows = list(ws.iter_rows(values_only=True))
                            headers = rows[0]
                            data_rows = rows[1:]

                            # حذف البيانات القديمة (اختياري)
                            c.execute(f"DELETE FROM {table_name}")

                            for row in data_rows:
                                placeholders = ",".join("?" for _ in row)
                                c.execute(f"INSERT INTO {table_name} ({','.join(columns)}) VALUES ({placeholders})", row)

                    conn.commit()
                    conn.close()
                    msg(page, tr("Success", lang_module.lang))
                    page.update()

            file_picker.on_result = on_file_selected
            file_picker.pick_files(allow_multiple=False, allowed_extensions=["xlsx"])

        except Exception as ex:
            print("Import Error:", ex)
            error(page, tr("Error", lang_module.lang))
            page.update()

    return Column(
    scroll=ScrollMode.AUTO,
    horizontal_alignment="center",
    controls=[
        Row([
            IconButton(Icons.ARROW_BACK, tooltip=tr("back"), on_click=lambda e: go_back())
        ]),
        Text(tr("data_export_import", lang_module.lang), size=26, weight="bold"),
        Divider(),
        Card(
            content=Container(
                padding=20,
                content=Column([
                    Text(tr("export_data"), size=20, weight="bold"),
                    ElevatedButton(tr("export_data"), icon=Icons.UPLOAD_FILE, on_click=export_data)
                ])
            )
        ),
        Card(
            content=Container(
                padding=20,
                content=Column([
                    Text(tr("import_data"), size=20, weight="bold"),
                    ElevatedButton(tr("import_data"), icon=Icons.DOWNLOAD, on_click=import_data)
                ])
            )
        )
    ]
)



def build_appbar(page, set_lang, go_to_main_menu):
    return AppBar(
        bgcolor=Colors.GREEN_900,
        title=Text(tr("app_name", lang_module.lang), color=Colors.WHITE),
        color=Colors.WHITE,
        center_title=True,
        actions=[
            PopupMenuButton(
                items=[
                    PopupMenuItem(text="اللغة العربية", on_click=lambda e: (set_lang("ar"), go_to_main_menu())),
                    PopupMenuItem(text="English (EN)", on_click=lambda e: (set_lang("en"), go_to_main_menu())),
                    PopupMenuItem(text="About", on_click=lambda e: about(e.page)),
                    PopupMenuItem(),
                    PopupMenuItem(text="Exit", on_click=lambda e: page.window.close())
                ]
            )
        ]
    )



# ========== التطبيق ==========
def main(page: Page):
    page.title = "تشخيص الأمراض"
    page.window_width = 400
    page.window_height = 700
    page.vertical_alignment = "center"
    page.horizontal_alignment = "center"
    page.scroll = ScrollMode.AUTO
    page.theme_mode = ThemeMode.LIGHT
    page.navigation_bar=CupertinoNavigationBar(
        bgcolor=Colors.GREEN_900,
        inactive_color=Colors.BLACK,
                )
    def set_lang(selected_lang):
        lang_module.lang = selected_lang
        page.appbar = build_appbar(page, set_lang, go_to_main_menu)
        page.update()

    def go_to_main_menu():
        page.clean()
        page.appbar = build_appbar(page, set_lang, go_to_main_menu)
        page.add(MainMenuScreen(page, lang_module.lang, go_to_screen))
        page.update()

    def go_to_screen(screen_name):
        page.clean()
        page.appbar = build_appbar(page, set_lang, go_to_main_menu)
        if screen_name == "diagnosis":
            page.add(DiagnosisScreen(page, go_back=go_to_main_menu))
        elif screen_name == "add_disease":
            page.add(AddDiseaseScreen(page, go_back=go_to_main_menu))
        elif screen_name == "manage_symptoms":
            page.add(ManageSymptomsScreen(page, go_back=go_to_main_menu))
        elif screen_name == "history":
            page.add(DiagnosisHistoryScreen(page, go_back=go_to_main_menu))
        elif screen_name == "data_analysis":
            page.add(DataAnalysisScreen(page, go_back=go_to_main_menu))
        elif screen_name == "data_export_import":
            page.add(ExportImportScreen(page, go_back=go_to_main_menu))
        else:
            go_to_main_menu()

    # أول تعيين
    page.appbar = build_appbar(page, set_lang, go_to_main_menu)
    page.add(SplashScreen(page, set_lang, go_to_main_menu))

# تشغيل التطبيق
app(target=main)
